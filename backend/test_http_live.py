import io
import httpx
from docx import Document
import fitz

def test_live_http_word_to_pdf():
    # 1. Create DOCX in-memory
    doc = Document()
    doc.add_heading("Uji Coba Live Endpoint Word ke PDF", level=0)
    doc.add_paragraph("Pengujian HTTP live server ChangeFilePDF dengan LibreOffice headless.")
    doc.add_page_break()
    doc.add_paragraph("Halaman 2 dari dokumen uji coba live.")
    
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # 2. Send POST request to live server
    url = "http://127.0.0.1:8000/api/conversion/word-to-pdf"
    files = {
        "file": ("dokumen_live.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }

    with httpx.Client(timeout=30.0) as client:
        response = client.post(url, files=files)
        
    assert response.status_code == 200, f"HTTP request failed: {response.text}"
    assert response.headers.get("content-type") == "application/pdf"
    
    # 3. Validate generated PDF
    pdf_doc = fitz.open(stream=response.content, filetype="pdf")
    assert len(pdf_doc) == 2, f"Expected 2 pages, got {len(pdf_doc)}"
    print("=== PENGUJIAN LIVE HTTP SERVER BERHASIL ===")
    print(f"Status Code: {response.status_code}")
    print(f"Content-Type: {response.headers.get('content-type')}")
    print(f"PDF Output Size: {len(response.content)} bytes")
    print(f"Total Halaman: {len(pdf_doc)} halaman")
    pdf_doc.close()

if __name__ == "__main__":
    test_live_http_word_to_pdf()
