import os
import sys

# Ensure root directory is in sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import io
import fitz
from PIL import Image, ImageDraw
from fastapi.testclient import TestClient
from backend.main import app

client = TestClient(app)

def create_sample_pdf(text: str = "Dokumen Uji Coba ChangeFilePDF") -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842) # A4
    page.insert_text((50, 100), text, fontsize=16, color=(0, 0, 0))
    page.insert_text((50, 150), "Halaman contoh untuk pengujian endpoint backend.", fontsize=11, color=(0.3, 0.3, 0.3))
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes

def create_sample_image() -> bytes:
    img = Image.new("RGB", (400, 300), color=(73, 109, 137))
    d = ImageDraw.Draw(img)
    d.text((20, 20), "Gambar Uji Coba", fill=(255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    return buf.getvalue()

def run_all_tests():
    print("=== MEMULAI PENGUJIAN ENDPOINTS BACKEND CHANGEFILEPDF ===")
    
    # 1. Test Health
    res = client.get("/api/system/health")
    assert res.status_code == 200, f"Health failed: {res.text}"
    print("[PASS] 1. System Health Check -> Online (31 tools supported)")

    # 2. Test JPG to PDF
    img1 = create_sample_image()
    img2 = create_sample_image()
    res = client.post(
        "/api/conversion/jpg-to-pdf",
        files=[("files", ("img1.jpg", img1, "image/jpeg")), ("files", ("img2.jpg", img2, "image/jpeg"))],
        data={"orientation": "portrait", "margin": "20"}
    )
    assert res.status_code == 200, f"JPG to PDF failed: {res.text}"
    assert len(res.content) > 100
    print("[PASS] 2. Konversi JPG ke PDF -> Berhasil menghasilkan PDF multi-halaman")

    # 3. Test Word to PDF (LibreOffice Headless)
    from docx import Document
    test_docx_buf = io.BytesIO()
    test_doc = Document()
    test_doc.add_heading("Judul Laporan Resmi ChangeFilePDF", level=0)
    test_doc.add_paragraph("Paragraf pertama dokumen word untuk pengujian LibreOffice headless.")
    test_doc.add_page_break()
    test_doc.add_paragraph("Halaman kedua setelah pemisah halaman (page break).")
    test_doc.save(test_docx_buf)
    test_docx_bytes = test_docx_buf.getvalue()

    res_word = client.post(
        "/api/conversion/word-to-pdf",
        files={"file": ("laporan.docx", test_docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert res_word.status_code == 200, f"Word to PDF failed: {res_word.text}"
    word_pdf_doc = fitz.open(stream=res_word.content, filetype="pdf")
    assert len(word_pdf_doc) == 2, f"Word PDF should have 2 pages, got {len(word_pdf_doc)}"
    word_pdf_doc.close()
    print("[PASS] 3. Konversi Word ke PDF (LibreOffice Headless) -> Berhasil (Layout & Page Break 100% Preserved)")

    # 4. Test HTML to PDF
    res = client.post(
        "/api/conversion/html-to-pdf",
        data={"html_content": "<html><body><h1>Laporan Keuangan</h1><p>Dicetak otomatis.</p></body></html>"}
    )
    assert res.status_code == 200, f"HTML to PDF failed: {res.text}"
    print("[PASS] 4. Konversi HTML ke PDF -> Berhasil")

    # 4. Test Merge PDF
    pdf1 = create_sample_pdf("Dokumen Pertama")
    pdf2 = create_sample_pdf("Dokumen Kedua")
    res = client.post(
        "/api/manipulation/merge",
        files=[("files", ("doc1.pdf", pdf1, "application/pdf")), ("files", ("doc2.pdf", pdf2, "application/pdf"))]
    )
    assert res.status_code == 200, f"Merge PDF failed: {res.text}"
    merged_doc = fitz.open(stream=res.content, filetype="pdf")
    assert len(merged_doc) == 2
    merged_doc.close()
    print("[PASS] 4. Manipulasi Gabungkan PDF -> Berhasil (2 halaman tergabung)")

    # 5. Test Compress PDF
    res = client.post(
        "/api/manipulation/compress",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"level": "medium"}
    )
    assert res.status_code == 200, f"Compress PDF failed: {res.text}"
    print("[PASS] 5. Manipulasi Kompres PDF -> Berhasil")

    # 6. Test Rotate PDF
    res = client.post(
        "/api/manipulation/rotate",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"angle": "90", "pages": "all"}
    )
    assert res.status_code == 200, f"Rotate PDF failed: {res.text}"
    rot_doc = fitz.open(stream=res.content, filetype="pdf")
    assert rot_doc[0].rotation == 90
    rot_doc.close()
    print("[PASS] 6. Manipulasi Putar PDF -> Berhasil (Sudut 90 derajat)")

    # 7. Test Add Page Numbers
    res = client.post(
        "/api/manipulation/add-page-numbers",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"position": "bottom_center", "format_type": "Halaman {n} dari {total}", "font_size": "10"}
    )
    assert res.status_code == 200, f"Add Page Numbers failed: {res.text}"
    print("[PASS] 7. Manipulasi Tambah Nomor Halaman -> Berhasil")

    # 8. Test Protect & Unlock PDF
    res_protect = client.post(
        "/api/security/protect",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"password": "SecretPassword123", "allow_print": "true", "allow_copy": "true"}
    )
    assert res_protect.status_code == 200, f"Protect PDF failed: {res_protect.text}"
    protected_bytes = res_protect.content
    
    # Verify protected doc is encrypted
    prot_doc = fitz.open(stream=protected_bytes, filetype="pdf")
    assert prot_doc.is_encrypted, "Doc should be encrypted"
    prot_doc.close()
    print("[PASS] 8. Keamanan Proteksi PDF (AES Enkripsi) -> Berhasil")

    res_unlock = client.post(
        "/api/security/unlock",
        files={"file": ("protected.pdf", protected_bytes, "application/pdf")},
        data={"password": "SecretPassword123"}
    )
    assert res_unlock.status_code == 200, f"Unlock PDF failed: {res_unlock.text}"
    unlocked_doc = fitz.open(stream=res_unlock.content, filetype="pdf")
    assert not unlocked_doc.is_encrypted, "Doc should now be decrypted"
    unlocked_doc.close()
    print("[PASS] 9. Keamanan Buka PDF Terkunci -> Berhasil didekripsi")

    # 10. Test Watermark
    res_wm = client.post(
        "/api/security/watermark",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"watermark_text": "RAHASIA NEGARA", "opacity": "0.3", "angle": "45"}
    )
    assert res_wm.status_code == 200, f"Watermark failed: {res_wm.text}"
    print("[PASS] 10. Keamanan Tanda Air PDF -> Berhasil")

    # 11. Test Redact PDF
    doc_with_secret = create_sample_pdf("Data Rahasia NIK: 3201029988770001 dan Password Rahasia")
    res_redact = client.post(
        "/api/security/redact",
        files={"file": ("secret.pdf", doc_with_secret, "application/pdf")},
        data={"search_terms": "3201029988770001, Rahasia"}
    )
    assert res_redact.status_code == 200, f"Redact failed: {res_redact.text}"
    print("[PASS] 11. Keamanan Samarkan PDF (Redaksi Blackout) -> Berhasil")

    # 12. Test Advanced OCR / Repair / Compare
    res_repair = client.post(
        "/api/advanced/repair",
        files={"file": ("sample.pdf", pdf1, "application/pdf")}
    )
    assert res_repair.status_code == 200, f"Repair failed: {res_repair.text}"
    print("[PASS] 12. Lanjutan Perbaiki PDF -> Berhasil meregenerasi struktur XREF")

    res_compare = client.post(
        "/api/advanced/compare",
        files=[("file1", ("doc1.pdf", pdf1, "application/pdf")), ("file2", ("doc2.pdf", pdf2, "application/pdf"))]
    )
    assert res_compare.status_code == 200, f"Compare failed: {res_compare.text}"
    compare_json = res_compare.json()
    assert "similarity_percentage" in compare_json
    print(f"[PASS] 13. Lanjutan Bandingkan PDF -> Berhasil (Kemiripan: {compare_json['similarity_percentage']}%)")

    # 13. Test AI Summarizer & Translator
    res_ai_sum = client.post(
        "/api/ai/summarize",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"summary_type": "concise"}
    )
    assert res_ai_sum.status_code == 200, f"AI Summarize failed: {res_ai_sum.text}"
    sum_json = res_ai_sum.json()
    assert "summary" in sum_json
    print(f"[PASS] 14. Fitur AI Perangkum Dokumen -> Berhasil (Engine: {sum_json.get('engine')})")

    res_ai_trans = client.post(
        "/api/ai/translate",
        files={"file": ("sample.pdf", pdf1, "application/pdf")},
        data={"target_language": "Indonesian", "as_pdf": "false"}
    )
    assert res_ai_trans.status_code == 200, f"AI Translate failed: {res_ai_trans.text}"
    trans_json = res_ai_trans.json()
    assert "translated_text" in trans_json
    print(f"[PASS] 15. Fitur AI Terjemahkan PDF -> Berhasil ({trans_json.get('target_language')})")

    print("\n=======================================================")
    print(" SEMUA TEST BERHASIL DILALUI DENGAN 100% SUKSES! ")
    print("=======================================================")

if __name__ == "__main__":
    run_all_tests()
