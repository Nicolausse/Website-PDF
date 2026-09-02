import os
import sys

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import fitz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.services.conversion_service import ConversionService
from backend.utils.helpers import create_temp_dir

def create_complex_sample_docx(file_path: str) -> str:
    """Create a rich DOCX with cover page, page break, styled table, and formatting."""
    doc = Document()
    
    # 1. Cover Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("LAPORAN RESMI & DOKUMEN PROPOSAL\nChangeFilePDF")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Dark Blue
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Dokumen Uji Coba Konversi Headless LibreOffice Berkualitas Tinggi")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    
    doc.add_page_break() # Page break 1 -> Page 2
    
    # 2. Section on Page 2
    h1 = doc.add_heading("1. Rincian Fitur & Keamanan", level=1)
    p2 = doc.add_paragraph("Berikut adalah tabel perbandingan performa konversi:")
    
    # Styled Table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ["Modul", "Status", "Akurasi Layout"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        
    data = [
        ["Word ke PDF", "LibreOffice Headless", "100% Sempurna"],
        ["Format Gambar & Cover", "Preserved", "High Resolution"]
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.cell(row_idx + 1, col_idx).text = text
            
    doc.add_page_break() # Page break 2 -> Page 3
    
    # 3. Section on Page 3
    doc.add_heading("2. Kesimpulan & Verifikasi", level=1)
    p3 = doc.add_paragraph("Halaman ini membuktikan bahwa page break dan pemisahan halaman berhasil dipertahankan.")
    
    doc.save(file_path)
    return file_path

def test_word_to_pdf_conversion():
    temp_dir = create_temp_dir()
    docx_path = os.path.join(temp_dir, "sample_complex.docx")
    pdf_path = os.path.join(temp_dir, "sample_complex.pdf")
    
    print("1. Membuat berkas .docx sampel dengan halaman sampul & page break...")
    create_complex_sample_docx(docx_path)
    assert os.path.exists(docx_path)
    print(f"   Berkas .docx dibuat di: {docx_path} ({os.path.getsize(docx_path)} bytes)")
    
    print("\n2. Mengonversi .docx ke .pdf menggunakan LibreOffice Headless...")
    ConversionService.word_to_pdf(docx_path, pdf_path)
    
    assert os.path.exists(pdf_path), "File PDF harus terbentuk"
    pdf_size = os.path.getsize(pdf_path)
    assert pdf_size > 0, "File PDF tidak boleh kosong"
    print(f"   Berkas PDF berhasil dihasilkan: {pdf_path} ({pdf_size} bytes)")
    
    print("\n3. Memvalidasi integritas struktur PDF hasil konversi...")
    doc = fitz.open(pdf_path)
    page_count = len(doc)
    print(f"   Total Halaman PDF: {page_count} halaman")
    assert page_count == 3, f"Harus memiliki tepat 3 halaman sesuai page break (Ditemukan: {page_count})"
    
    # Inspect text on page 1, 2, 3
    page1_text = doc[0].get_text("text")
    page2_text = doc[1].get_text("text")
    page3_text = doc[2].get_text("text")
    
    assert "ChangeFilePDF" in page1_text, "Halaman 1 harus berisi judul sampul"
    assert "Rincian Fitur" in page2_text or "Modul" in page2_text, "Halaman 2 harus berisi konten section 1"
    assert "Kesimpulan" in page3_text, "Halaman 3 harus berisi konten section 2"
    
    doc.close()
    print("\n[BERHASIL] Uji coba Word ke PDF via LibreOffice Headless 100% Lulus Sempurna!")

if __name__ == "__main__":
    test_word_to_pdf_conversion()
